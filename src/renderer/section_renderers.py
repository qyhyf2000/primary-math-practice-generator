"""各题型渲染函数"""
import json
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .layout import (
    add_paragraph, set_cjk_font, set_cell_border,
    cell_add_run, set_table_borders,
)
from .graphics_renderer import (
    render_angle_question, render_count_angles_question,
    render_grid_count_question, render_angle_drawing,
    render_clock_question, render_clock_time_question,
    render_cube_stack_question, render_cube_view_question,
    render_shape_judge_question, render_shape_classify_question,
    render_tangram_question, render_parallelogram_question,
)
from ..question_bank.models import ExamSection


FONT = "宋体"
HEADING_FONT = "楷体"
BODY_SIZE = 12
HEADING_SIZE = 12


def _parse_graphic_info(q) -> dict:
    """从题目的 tags 中解析图形渲染信息。

    graphic 数据以 'graphic:' 开头，后面跟 JSON。
    由于 JSON 内部含逗号，不能简单按逗号 split。
    """
    if not q.tags:
        return None
    # 找到 graphic: 的位置
    idx = q.tags.find("graphic:")
    if idx < 0:
        return None
    # 提取 graphic: 之后的所有内容作为 JSON
    candidate = q.tags[idx + len("graphic:"):].strip()
    if not candidate.startswith("{"):
        return None

    # 优先直接解析（graphic JSON 通常是 tags 末尾的内容）
    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        pass

    # 回退：手动数大括号截取 JSON（兼容 tags 中 graphic 后有额外内容的情况）
    depth = 0
    end = 0
    for i, ch in enumerate(candidate):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i + 1
                break
    if end == 0:
        return None
    try:
        return json.loads(candidate[:end])
    except json.JSONDecodeError:
        return None


def _render_graphic_question(doc, idx: int, q):
    """根据 graphic 信息渲染图形题"""
    info = _parse_graphic_info(q)
    if not info:
        # 无图形信息，回退到普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{idx + 1}. {q.content}")
        set_cjk_font(run, FONT, BODY_SIZE)
        return

    gtype = info.get("type", "")

    if gtype == "angle_identify":
        angles = info.get("angles", [])
        render_angle_question(doc, idx + 1, q.content, angles)

    elif gtype == "angle_judge":
        shapes = info.get("shapes", [])
        render_shape_judge_question(doc, idx + 1, q.content, shapes)

    elif gtype == "count_angles":
        shapes = info.get("shapes", [])
        render_count_angles_question(doc, idx + 1, q.content, shapes)

    elif gtype == "grid_count":
        render_grid_count_question(
            doc, idx + 1, q.content,
            rows=info.get("rows", 2),
            cols=info.get("cols", 2),
        )

    elif gtype == "draw_grid":
        render_grid_count_question(
            doc, idx + 1, q.content,
            rows=info.get("rows", 4),
            cols=info.get("cols", 6),
        )

    elif gtype == "draw_angle":
        render_angle_drawing(doc, idx + 1, q.content)

    elif gtype == "clock":
        clocks = info.get("clocks", [])
        render_clock_question(doc, idx + 1, q.content, clocks)

    elif gtype == "clock_time":
        times = [(t["time"], t["label"]) for t in info.get("times", [])]
        render_clock_time_question(doc, idx + 1, q.content, times)

    elif gtype == "cube_stack":
        grid = info.get("grid", [[1]])
        render_cube_stack_question(doc, idx + 1, q.content, grid)

    elif gtype == "cube_view":
        render_cube_view_question(
            doc, idx + 1, q.content,
            front_view=info.get("front", [1]),
            side_view=info.get("side", [1]),
        )

    elif gtype == "shape_classify":
        shapes = info.get("shapes", [])
        render_shape_classify_question(doc, idx + 1, q.content, shapes)

    elif gtype == "tangram":
        grid = info.get("grid", [["r"]])
        pieces = info.get("pieces", [])
        render_tangram_question(doc, idx + 1, q.content, grid, pieces)

    elif gtype == "parallelogram":
        render_parallelogram_question(doc, idx + 1, q.content)

    else:
        # 未知图形类型，回退
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{idx + 1}. {q.content}")
        set_cjk_font(run, FONT, BODY_SIZE)


def render_section_heading(doc, section: ExamSection):
    """渲染题型标题，如：一、口算题（每题1分，共10分）"""
    title_text = f"{section.title}（每题{section.score_per_question}分，共{section.total_score}分）"
    add_paragraph(doc, title_text, HEADING_FONT, HEADING_SIZE,
                  bold=True, space_after_pt=4)


def render_oral_calc(doc, section: ExamSection):
    """口算题 — 表格布局，5列×N行"""
    questions = section.questions
    cols = 5
    rows = (len(questions) + cols - 1) // cols

    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, val="single", sz="4")

    for idx, q in enumerate(questions):
        r = idx // cols
        c = idx % cols
        cell = table.cell(r, c)
        cell_add_run(cell, f"{idx + 1}. {q.content}", FONT, BODY_SIZE)

    doc.add_paragraph()


def render_fill_blank(doc, section: ExamSection):
    """填空题 — 段落形式，图形题用专用渲染器"""
    for idx, q in enumerate(section.questions):
        if "图形" in (q.tags or ""):
            _render_graphic_question(doc, idx, q)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"{idx + 1}. {q.content}")
            set_cjk_font(run, FONT, BODY_SIZE)


def render_choice(doc, section: ExamSection):
    """选择题 — 每题一段，选项横向排列，图形题用专用渲染器"""
    for idx, q in enumerate(section.questions):
        if "图形" in (q.tags or ""):
            # 先用图形渲染器画图形
            _render_graphic_question(doc, idx, q)
            # 再渲染选项
            options = q.get_options_list()
            if options:
                opt_p = doc.add_paragraph()
                opt_p.paragraph_format.space_after = Pt(4)
                opt_p.paragraph_format.left_indent = Cm(0.8)
                opt_text = "     ".join(options)
                run = opt_p.add_run(opt_text)
                set_cjk_font(run, FONT, BODY_SIZE)
        else:
            # 题干
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{idx + 1}. {q.content}")
            set_cjk_font(run, FONT, BODY_SIZE)

            # 选项
            options = q.get_options_list()
            if options:
                opt_p = doc.add_paragraph()
                opt_p.paragraph_format.space_after = Pt(4)
                opt_p.paragraph_format.left_indent = Cm(0.8)
                opt_text = "     ".join(options)
                run = opt_p.add_run(opt_text)
                set_cjk_font(run, FONT, BODY_SIZE)


def render_vertical_calc(doc, section: ExamSection):
    """竖式/脱式计算 — 每题后留空白书写区"""
    for idx, q in enumerate(section.questions):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{idx + 1}. {q.content}")
        set_cjk_font(run, FONT, BODY_SIZE)

        # 空白书写区（用空行模拟）
        for _ in range(4):
            blank_p = doc.add_paragraph()
            blank_p.paragraph_format.space_before = Pt(0)
            blank_p.paragraph_format.space_after = Pt(0)
            # 添加一个底边框行作为书写线
            run = blank_p.add_run(" " * 60)
            set_cjk_font(run, FONT, BODY_SIZE)


def render_word_problem(doc, section: ExamSection):
    """解决问题（应用题）— 每题后留列式和答的空白"""
    for idx, q in enumerate(section.questions):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{idx + 1}. {q.content}")
        set_cjk_font(run, FONT, BODY_SIZE)

        # 列式区域
        formula_p = doc.add_paragraph()
        formula_p.paragraph_format.space_before = Pt(8)
        run = formula_p.add_run("列式：")
        set_cjk_font(run, FONT, BODY_SIZE, bold=True)

        for _ in range(3):
            blank_p = doc.add_paragraph()
            blank_p.paragraph_format.space_before = Pt(0)
            blank_p.paragraph_format.space_after = Pt(0)
            run = blank_p.add_run(" " * 60)
            set_cjk_font(run, FONT, BODY_SIZE)

        # 答区域
        answer_p = doc.add_paragraph()
        answer_p.paragraph_format.space_before = Pt(8)
        answer_p.paragraph_format.space_after = Pt(6)
        run = answer_p.add_run("答：")
        set_cjk_font(run, FONT, BODY_SIZE, bold=True)
        run = answer_p.add_run("_" * 50)
        set_cjk_font(run, FONT, BODY_SIZE)


RENDERER_MAP = {
    "oral_calc": render_oral_calc,
    "fill_blank": render_fill_blank,
    "choice": render_choice,
    "vertical_calc": render_vertical_calc,
    "word_problem": render_word_problem,
}
