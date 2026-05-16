"""页面布局和字体工具"""
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_page(doc, rendering_cfg):
    """设置A4页面和页边距"""
    page = rendering_cfg["page"]
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(page["margin_top_cm"])
    section.bottom_margin = Cm(page["margin_bottom_cm"])
    section.left_margin = Cm(page["margin_left_cm"])
    section.right_margin = Cm(page["margin_right_cm"])


def set_cjk_font(run, font_name: str, size_pt: float = None,
                 bold: bool = False):
    """设置中文字体（同时设置西文和东亚字体）"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold


def add_paragraph(doc, text: str, font_name: str, size_pt: float,
                  bold: bool = False, alignment=None, space_after_pt: float = 6):
    """添加一个带中文字体的段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    set_cjk_font(run, font_name, size_pt, bold)
    return p


def set_cell_border(cell, **kwargs):
    """设置单元格边框。

    用法：set_cell_border(cell, top="none", bottom="single", left="none", right="none")
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")

    for edge, val in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val)
        if val != "none":
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
        tcBorders.append(element)

    tcPr.append(tcBorders)


def set_cell_shading(cell, color: str):
    """设置单元格背景色"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def set_table_borders(table, val: str = "single", sz: str = "4"):
    """设置整个表格的边框"""
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tblBorders.append(element)
    tblPr.append(tblBorders)


def cell_add_run(cell, text: str, font_name: str = "宋体",
                 size_pt: float = 12, bold: bool = False):
    """向单元格添加带中文字体的文本"""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_cjk_font(run, font_name, size_pt, bold)
    return run


def create_blank_cell(cell, width_cm: float = 2.0):
    """创建带底边框的空白单元格（模拟下划线填空）"""
    cell.width = Cm(width_cm)
    set_cell_border(cell, top="none", bottom="single", left="none", right="none")
    cell_add_run(cell, "　　", size_pt=12)
