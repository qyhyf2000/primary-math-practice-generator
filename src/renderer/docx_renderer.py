"""Word文档主渲染器 - 组装标题、考生信息、各题型区域"""
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..config.config_manager import ConfigManager
from ..question_bank.models import Exam
from .layout import (
    setup_page, add_paragraph, set_cjk_font, set_cell_border,
    cell_add_run, set_table_borders,
)
from .section_renderers import (
    render_section_heading, RENDERER_MAP,
    FONT, BODY_SIZE,
)


class DocxRenderer:
    def __init__(self, config: ConfigManager):
        self.config = config
        self.rendering_cfg = config.get_rendering_config()
        self.fonts = self.rendering_cfg["fonts"]

    def render(self, exam: Exam, output_path: str) -> str:
        """渲染试卷并保存为 .docx 文件"""
        doc = Document()
        setup_page(doc, self.rendering_cfg)

        self._render_title(doc, exam.title)
        self._render_student_info(doc)
        self._render_separator(doc)

        for section in exam.sections:
            render_section_heading(doc, section)
            render_fn = RENDERER_MAP.get(section.section_id)
            if render_fn:
                render_fn(doc, section)
            else:
                # 兜底：段落形式
                self._render_fallback(doc, section)

        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path

    def _render_title(self, doc, title: str):
        """试卷标题 — 黑体居中加粗"""
        title_cfg = self.rendering_cfg["title"]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        set_cjk_font(run, self.fonts["title"],
                     title_cfg["font_size_pt"], bold=title_cfg["bold"])

    def _render_student_info(self, doc):
        """考生信息栏 — 无边框表格 + 底边框空白模拟下划线"""
        info_cfg = self.rendering_cfg["student_info"]
        fields = info_cfg["fields"]
        size = info_cfg["font_size_pt"]

        table = doc.add_table(rows=1, cols=len(fields) * 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, field in enumerate(fields):
            label_cell = table.cell(0, i * 2)
            blank_cell = table.cell(0, i * 2 + 1)

            # 标签：移除所有边框
            set_cell_border(label_cell, top="none", bottom="none",
                          left="none", right="none")
            label_cell.width = Cm(1.5)
            cell_add_run(label_cell, f"{field}：", self.fonts["body"], size)

            # 空白下划线
            blank_cell.width = Cm(3.5)
            set_cell_border(blank_cell, top="none", bottom="single",
                          left="none", right="none")
            cell_add_run(blank_cell, "　　　　　　", self.fonts["body"], size)

        doc.add_paragraph()

    def _render_separator(self, doc):
        """分隔线"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        # 用边框线分隔
        pPr = p._element.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def render_with_answer(self, exam: Exam, output_path: str) -> str:
        """渲染试卷 + 答案页（家长版）"""
        doc = Document()
        setup_page(doc, self.rendering_cfg)

        # === 学生版（前几页） ===
        self._render_title(doc, exam.title + "（学生版）")
        self._render_student_info(doc)
        self._render_separator(doc)

        for section in exam.sections:
            render_section_heading(doc, section)
            render_fn = RENDERER_MAP.get(section.section_id)
            if render_fn:
                render_fn(doc, section)
            else:
                self._render_fallback(doc, section)

        # === 分页：答案页 ===
        doc.add_page_break()

        self._render_title(doc, exam.title + "（参考答案）")
        p_note = doc.add_paragraph()
        p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_note = p_note.add_run("—— 家长批改专用，请勿给学生 ——")
        set_cjk_font(run_note, self.fonts["body"], 10)

        for section in exam.sections:
            # 题型标题
            render_section_heading(doc, section)
            for idx, q in enumerate(section.questions):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                ans_text = q.answer or ""
                if q.options and q.section == "choice":
                    ans_text = f"{q.answer}"
                run = p.add_run(f"{idx + 1}. {ans_text}")
                set_cjk_font(run, self.fonts["body"], 11)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path

    def _render_fallback(self, doc, section):
        """兜底渲染 — 段落形式逐题输出"""
        for idx, q in enumerate(section.questions):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"{idx + 1}. {q.content}")
            set_cjk_font(run, FONT, BODY_SIZE)
