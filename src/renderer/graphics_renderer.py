"""
图形题渲染器 — 用表格和Unicode字符在Word文档中绘制小学数学几何图形

支持的类型：
- angle: 角的识别（直角/锐角/钝角）
- count_angles: 数图形中的角
- count_rects: 数长方形/正方形个数
- shape_identify: 图形识别（长方形/正方形/平行四边形）
- grid_count: 网格中的图形计数
- clock: 钟面读时（PIL 绘制圆形钟面）
"""
import math
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont

from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .layout import (
    set_cjk_font, set_cell_border, cell_add_run,
    set_table_borders, set_cell_shading,
)

FONT = "宋体"
BODY_SIZE = 12


def _get_pil_font(size: int = 16):
    """加载中文字体，找不到系统字体时回退默认"""
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for fp in candidates:
        try:
            return ImageFont.truetype(fp, size)
        except (OSError, IOError):
            continue
    return ImageFont.load_default()


def _draw_angle_image(angle_type: str, size: int = 120) -> BytesIO:
    """
    用 PIL 绘制单个角的图形。

    angle_type: "直角" / "锐角" / "钝角"
    size: 图片像素尺寸
    """
    import math as _math
    img = Image.new('RGB', (size, size), 'white')
    draw = ImageDraw.Draw(img)

    cx, cy = size // 2, size - 20  # 顶点在底部中央
    ray_len = size - 50
    lw = 3  # 线宽

    if angle_type == "直角":
        angle_deg = 90
    elif angle_type == "锐角":
        angle_deg = 50
    else:  # 钝角
        angle_deg = 130

    # 水平射线（向右）
    end_h = (cx + ray_len, cy)
    # 另一条射线（按角度逆时针旋转）
    rad = _math.radians(angle_deg)
    end_a = (cx - int(ray_len * _math.cos(rad)),
             cy - int(ray_len * _math.sin(rad)))

    draw.line([end_a, (cx, cy), end_h], fill=(60, 60, 60), width=lw)

    # 画角度弧线
    arc_r = 22
    if angle_type == "直角":
        # 直角标记：小正方形
        sq = 14
        draw.rectangle([cx, cy - sq, cx + sq, cy], outline=(60, 60, 60), width=2)
    else:
        # 弧线
        start_angle = 0
        end_angle = -angle_deg
        bbox = [cx - arc_r, cy - arc_r, cx + arc_r, cy + arc_r]
        draw.arc(bbox, start=end_angle, end=0, fill=(60, 60, 60), width=2)

    # 画顶点小圆点
    draw.ellipse([cx - 4, cy - 4, cx + 4, cy + 4], fill=(60, 60, 60))

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def render_angle_question(doc, q_num: int, q_content: str, angles: list):
    """
    渲染角的识别题 —— 用 PIL 绘制清晰的角图形。

    angles: [{"symbol": "╲", "label": "第1个"}, ...]
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    # 用表格排列角图形
    table = doc.add_table(rows=2, cols=len(angles))
    table.alignment = 1
    set_table_borders(table, val="none")

    # 映射 symbol → 角类型
    SYMBOL_TO_TYPE = {
        "┌": "直角", "∟": "直角",
        "∠": "锐角", "╱": "锐角",
        "╲": "钝角", "╱ ": "钝角",
    }

    for i, angle_info in enumerate(angles):
        symbol = angle_info.get("symbol", "∠")
        label = angle_info.get("label", f"第{i+1}个")
        angle_type = SYMBOL_TO_TYPE.get(symbol, "锐角")

        # 图片行
        cell_img = table.cell(0, i)
        cell_img.width = Cm(3.0)
        p_img = cell_img.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_buf = _draw_angle_image(angle_type)
        run_img = p_img.add_run()
        run_img.add_picture(img_buf, width=Cm(2.5))

        # 标签行
        cell_lbl = table.cell(1, i)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_lbl = p_lbl.add_run(f"({label})")
        set_cjk_font(run_lbl, FONT, 10)

    doc.add_paragraph()


def _draw_polygon_shape(sides: int, size: int = 100) -> BytesIO:
    """用 PIL 绘制正多边形（用于数角题）"""
    img = Image.new('RGB', (size, size), 'white')
    draw = ImageDraw.Draw(img)
    cx, cy = size // 2, size // 2
    r = size // 2 - 8
    lw = 3

    # 计算顶点
    pts = []
    for i in range(sides):
        angle = -math.pi / 2 + 2 * math.pi * i / sides
        x = cx + int(r * math.cos(angle))
        y = cy + int(r * math.sin(angle))
        pts.append((x, y))

    draw.polygon(pts, outline=(60, 60, 60), width=lw)
    # 顶点小圆点
    for x, y in pts:
        draw.ellipse([x - 3, y - 3, x + 3, y + 3], fill=(60, 60, 60))

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def _draw_geometric_shape(shape_type: str, size: int = 100) -> BytesIO:
    """用 PIL 绘制几何图形（长方形/正方形/平行四边形/三角形）"""
    img = Image.new('RGB', (size, size), 'white')
    draw = ImageDraw.Draw(img)
    lw = 3
    m = 15  # 边距

    if shape_type == "square":
        draw.rectangle([m, m, size - m, size - m], outline=(60, 60, 60), width=lw)
    elif shape_type == "rectangle":
        draw.rectangle([m, m + 10, size - m, size - m - 10], outline=(60, 60, 60), width=lw)
    elif shape_type == "parallelogram":
        skew = 18
        pts = [(m + skew, m + 10), (size - m + skew, m + 10),
               (size - m, size - m - 10), (m, size - m - 10)]
        draw.polygon(pts, outline=(60, 60, 60), width=lw)
    elif shape_type == "triangle":
        pts = [(size // 2, m), (size - m, size - m), (m, size - m)]
        draw.polygon(pts, outline=(60, 60, 60), width=lw)
    elif shape_type == "pentagon":
        pts = _get_polygon_pts(size, 5)
        draw.polygon(pts, outline=(60, 60, 60), width=lw)
    elif shape_type == "hexagon":
        pts = _get_polygon_pts(size, 6)
        draw.polygon(pts, outline=(60, 60, 60), width=lw)
    else:
        draw.rectangle([m, m, size - m, size - m], outline=(60, 60, 60), width=lw)

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def _get_polygon_pts(size: int, sides: int) -> list:
    cx = cy = size // 2
    r = size // 2 - 10
    pts = []
    for i in range(sides):
        angle = -math.pi / 2 + 2 * math.pi * i / sides
        pts.append((cx + int(r * math.cos(angle)), cy + int(r * math.sin(angle))))
    return pts


# symbol → 角类型映射（复用于 render_shape_judge）
_SYMBOL_TO_ANGLE = {
    "∠": "锐角", "╱": "锐角",
    "∟": "直角", "┌": "直角",
    "╲": "钝角",
    "┐": "直角",
}
# symbol → 几何图形类型映射
_SYMBOL_TO_SHAPE = {
    "□": "square", "▭": "rectangle",
    "▱": "parallelogram", "△": "triangle",
    "⬠": "pentagon", "⬡": "hexagon",
    "○": "circle",
}


def render_count_angles_question(doc, q_num: int, q_content: str, shapes: list):
    """
    渲染数角题 —— PIL 绘制正多边形，学生数角的个数。

    shapes: [("△", "三角形"), ("□", "正方形"), ...]
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    cols = len(shapes)
    table = doc.add_table(rows=2, cols=cols)
    table.alignment = 1
    set_table_borders(table, val="none")

    SIDES_MAP = {"△": 3, "□": 4, "⬠": 5, "⬡": 6, "▭": 4, "▱": 4}

    for shape_info in shapes:
        symbol = shape_info.get("symbol", "△").strip() if isinstance(shape_info, dict) else shape_info[0]
        label = shape_info.get("label", "") if isinstance(shape_info, dict) else (shape_info[1] if len(shape_info) > 1 else "")
        i = shapes.index(shape_info)

        # 图形图片
        cell_img = table.cell(0, i)
        cell_img.width = Cm(3.0)
        p1 = cell_img.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sides = SIDES_MAP.get(symbol, 4)
        img_buf = _draw_polygon_shape(sides)
        run_img = p1.add_run()
        run_img.add_picture(img_buf, width=Cm(2.2))

        # 标签 + 作答
        cell_ans = table.cell(1, i)
        p2 = cell_ans.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"{label}\n(    )个角")
        set_cjk_font(run2, FONT, 10)

    doc.add_paragraph()


def render_shape_judge_question(doc, q_num: int, q_content: str, shapes: list):
    """
    渲染角度/图形判断题 —— 用 PIL 绘制角的图形或几何形状。

    shapes: [{"symbol": "∠", "label": "①"}, ...]
    符号中含 "┐┌" 等多字符时，绘制两个分离的线段表示"不是角"。
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    cols = len(shapes)
    table = doc.add_table(rows=2, cols=cols)
    table.alignment = 1
    set_table_borders(table, val="none")

    for i, shape_info in enumerate(shapes):
        symbol = shape_info.get("symbol", "∠")
        label = shape_info.get("label", f"①")

        # 图片
        cell_img = table.cell(0, i)
        cell_img.width = Cm(3.0)
        p1 = cell_img.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if len(symbol) >= 2 and symbol != symbol[0] * len(symbol):
            # 多字符符号（如 "┐┌"）= 不是角，画两条分离线段
            img_buf = _draw_not_an_angle()
        elif symbol in _SYMBOL_TO_ANGLE:
            img_buf = _draw_angle_image(_SYMBOL_TO_ANGLE[symbol])
        else:
            img_buf = _draw_angle_image("锐角")

        run_img = p1.add_run()
        run_img.add_picture(img_buf, width=Cm(2.2))

        # 标签 + 作答
        cell_ans = table.cell(1, i)
        p2 = cell_ans.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"{label} (    )")
        set_cjk_font(run2, FONT, 10)

    doc.add_paragraph()


def _draw_not_an_angle(size: int = 120) -> BytesIO:
    """绘制两条分离的线段，表示'不是角'"""
    img = Image.new('RGB', (size, size), 'white')
    draw = ImageDraw.Draw(img)
    lw = 3
    # 水平线段
    draw.line([(20, 30), (size - 60, 30)], fill=(60, 60, 60), width=lw)
    # 垂直线段（不相连）
    draw.line([(size - 35, 20), (size - 35, size - 40)], fill=(60, 60, 60), width=lw)
    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def render_shape_classify_question(doc, q_num: int, q_content: str, shapes: list):
    """
    渲染图形分类/命名题 —— PIL 绘制清晰的几何图形。

    shapes: [{"symbol": "▭", "label": "①"}, ...]
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    cols = len(shapes)
    table = doc.add_table(rows=2, cols=cols)
    table.alignment = 1
    set_table_borders(table, val="none")

    for i, shape_info in enumerate(shapes):
        symbol = shape_info.get("symbol", "□").strip()
        label = shape_info.get("label", "")

        # 图片
        cell_img = table.cell(0, i)
        cell_img.width = Cm(3.0)
        p1 = cell_img.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape_type = _SYMBOL_TO_SHAPE.get(symbol, "square")
        img_buf = _draw_geometric_shape(shape_type)
        run_img = p1.add_run()
        run_img.add_picture(img_buf, width=Cm(2.2))

        # 标签 + 作答
        cell_ans = table.cell(1, i)
        p2 = cell_ans.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl = label if label else f"图形{i + 1}"
        run2 = p2.add_run(f"{lbl}: __________")
        set_cjk_font(run2, FONT, 10)

    doc.add_paragraph()


def _draw_grid_image(rows: int, cols: int, cell_size: int = 60) -> BytesIO:
    """
    用 PIL 绘制清晰的方格网格图。

    cell_size: 每格像素大小
    """
    lw = 2  # 线宽
    w = cols * cell_size + lw
    h = rows * cell_size + lw

    img = Image.new('RGB', (w, h), 'white')
    draw = ImageDraw.Draw(img)

    # 画横线
    for r in range(rows + 1):
        y = r * cell_size
        draw.line([(0, y), (w, y)], fill=(40, 40, 40), width=lw)

    # 画竖线
    for c in range(cols + 1):
        x = c * cell_size
        draw.line([(x, 0), (x, h)], fill=(40, 40, 40), width=lw)

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def render_grid_count_question(doc, q_num: int, q_content: str,
                               rows: int, cols: int, answer: str = "",
                               description: str = ""):
    """
    渲染网格图形计数题 —— 用 PIL 绘制清晰网格。

    绘制一个 rows x cols 的完整网格，下方留空作答。
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    # 生成网格图片并嵌入
    img_buf = _draw_grid_image(rows, cols)
    if img_buf:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_buf, width=Cm(6))

    # 题目内容已包含填空括号，下方仅放网格图即可
    doc.add_paragraph()


def _draw_tangram_house(pieces: list = None) -> BytesIO:
    """
    用 PIL 绘制七巧板房子 —— 用真实的三角形、正方形、平行四边形拼块。

    7 块七巧板按经典房子图案排列：
    - 2 个大三角形（红） → 屋顶
    - 1 个中三角形（蓝） → 屋顶下三角区
    - 1 个正方形（黄）   → 墙体
    - 2 个小三角形（绿） → 左侧装饰
    - 1 个平行四边形（紫） → 门
    """
    img_w, img_h = 360, 290
    img = Image.new('RGB', (img_w, img_h), 'white')
    draw = ImageDraw.Draw(img)

    color_map = {
        'r': (235, 95, 95), 'b': (70, 165, 200), 'g': (115, 195, 135),
        'y': (245, 215, 80), 'p': (195, 140, 205),
        'o': (250, 160, 75), 'c': (125, 200, 230),
    }
    name_map = {
        'r': '大三角形', 'b': '中三角形', 'g': '小三角形',
        'y': '正方形', 'p': '平行四边形',
        'o': '大三角形', 'c': '小三角形',
    }
    lw = 3

    # === 7 块拼板的多边形顶点 ===
    # 屋顶左：大三角形（红）
    roof_left = [(180, 30), (45, 140), (180, 140)]
    # 屋顶右：大三角形（红）
    roof_right = [(180, 30), (315, 140), (180, 140)]
    # 屋顶下方：中三角形（蓝）
    mid_tri = [(180, 140), (250, 140), (180, 200)]
    # 墙体：正方形（黄）
    square_body = [(100, 140), (180, 140), (180, 220), (100, 220)]
    # 左侧上：小三角形（绿）
    small_top = [(45, 140), (100, 140), (45, 185)]
    # 左侧下：小三角形（绿）
    small_bot = [(45, 185), (100, 220), (45, 250)]
    # 门：平行四边形（紫）
    para_door = [(100, 220), (180, 220), (200, 260), (120, 260)]

    # 颜色码顺序对应 pieces 列表
    piece_shapes = []
    if pieces:
        for name, code in pieces:
            piece_shapes.append(code)
    else:
        piece_shapes = ['r', 'r', 'b', 'y', 'g', 'g', 'p']

    all_verts = [roof_left, roof_right, mid_tri, square_body,
                 small_top, small_bot, para_door]

    for i, verts in enumerate(all_verts):
        code = piece_shapes[i] if i < len(piece_shapes) else 'g'
        color = color_map.get(code, (200, 200, 200))
        draw.polygon(verts, fill=color, outline=(50, 50, 50), width=lw)

        # 顶点小圆点
        for x, y in verts:
            draw.ellipse([x - 3, y - 3, x + 3, y + 3], fill=(50, 50, 50))

    # 图例
    legend_y = 260
    legend_x = 20
    seen = set()
    if pieces:
        for name, code in pieces:
            if code in seen:
                continue
            seen.add(code)
            color = color_map.get(code, (200, 200, 200))
            draw.rectangle(
                [legend_x, legend_y, legend_x + 16, legend_y + 14],
                fill=color, outline=(50, 50, 50), width=1
            )
            shape_name = name_map.get(code, name)
            draw.text((legend_x + 22, legend_y - 1), shape_name, fill=(50, 50, 50))
            legend_x += 100

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def render_tangram_question(doc, q_num: int, q_content: str,
                            grid: list = None, pieces: list = None):
    """
    渲染七巧板拼图题 —— PIL 绘制真正的三角形/正方形/平行四边形拼块。
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    img_buf = _draw_tangram_house(pieces)
    if img_buf:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_buf, width=Cm(10))

    doc.add_paragraph()


def _draw_parallelogram_comparison() -> BytesIO:
    """
    用 PIL 绘制长方形与平行四边形对比图。

    左侧为长方形（标直角），右侧为拉成的平行四边形（标锐角/钝角），
    中间加箭头表示拉伸变形。尺寸相同，对比边长不变 vs 角度变了。
    """
    rect_w, rect_h = 160, 100   # 长方形宽高
    skew = 35                    # 平行四边形偏移量
    pad = 40
    gap = 60                     # 两图间距

    img_w = rect_w * 2 + gap + pad * 2
    img_h = rect_h + pad * 3 + 30

    img = Image.new('RGB', (img_w, img_h), 'white')
    draw = ImageDraw.Draw(img)
    lw = 3  # 线宽

    # 颜色
    color_line = (50, 50, 50)
    color_label = (80, 80, 80)

    # === 左侧：长方形 ===
    rx1 = pad
    ry1 = pad + 20
    rx2 = rx1 + rect_w
    ry2 = ry1 + rect_h

    # 长方形边框
    draw.rectangle([rx1, ry1, rx2, ry2], outline=color_line, width=lw)

    # 四个直角标记
    sq = 12
    corners_rect = [
        (rx1, ry1, "右下"),   # 左上角
        (rx2, ry1, "左下"),   # 右上角
        (rx2, ry2, "左上"),   # 右下角
        (rx1, ry2, "右上"),   # 左下角
    ]
    for cx, cy, _ in corners_rect:
        if cx == rx1 and cy == ry1:  # 左上
            draw.rectangle([cx, cy, cx + sq, cy + sq], outline=color_line, width=2)
        elif cx == rx2 and cy == ry1:  # 右上
            draw.rectangle([cx - sq, cy, cx, cy + sq], outline=color_line, width=2)
        elif cx == rx2 and cy == ry2:  # 右下
            draw.rectangle([cx - sq, cy - sq, cx, cy], outline=color_line, width=2)
        elif cx == rx1 and cy == ry2:  # 左下
            draw.rectangle([cx, cy - sq, cx + sq, cy], outline=color_line, width=2)

    # 边长标注
    draw.text((rx1 + rect_w // 2 - 20, ry2 + 8), "底边不变", fill=color_label)

    # === 箭头（中间） ===
    arrow_x = rx2 + 15
    arrow_y = ry1 + rect_h // 2
    draw.line([(arrow_x, arrow_y), (arrow_x + gap - 15, arrow_y)],
              fill=color_line, width=2)
    # 箭头尖
    ax = arrow_x + gap - 15
    draw.polygon([(ax, arrow_y), (ax - 10, arrow_y - 6), (ax - 10, arrow_y + 6)],
                 fill=color_line)
    draw.text((arrow_x + 5, arrow_y - 16), "拉", fill=color_label)

    # === 右侧：平行四边形 ===
    px1 = rx2 + gap
    py1 = ry1 + skew // 2
    px2 = px1 + rect_w
    py2 = py1 + rect_h

    # 平行四边形顶点（上边水平，下边右移 skew）
    para_pts = [
        (px1, py1),           # 左上
        (px2, py1),           # 右上
        (px2 + skew, py2),    # 右下
        (px1 + skew, py2),    # 左下
    ]
    draw.polygon(para_pts, outline=color_line, width=lw)

    # 角标记：锐角在左上和右下，钝角在右上和左下
    arc_r = 15
    # 左上角（锐角）：上边和左边夹角
    draw.arc([px1 - arc_r//2, py1 - arc_r//2, px1 + arc_r, py1 + arc_r],
             start=0, end=60, fill=color_line, width=2)
    # 右下角（锐角）
    brx, bry = px2 + skew, py2
    draw.arc([brx - arc_r, bry - arc_r, brx + arc_r//2, bry + arc_r//2],
             start=180, end=240, fill=color_line, width=2)
    # 右上角（钝角）
    draw.arc([px2 - arc_r, py1 - arc_r//2, px2 + arc_r//2, py1 + arc_r],
             start=120, end=180, fill=color_line, width=2)
    # 左下角（钝角）
    blx, bly = px1 + skew, py2
    draw.arc([blx - arc_r//2, bly - arc_r, blx + arc_r, bly + arc_r//2],
             start=300, end=0, fill=color_line, width=2)

    # 虚线提示高度（长方形高度 = 平行四边形高度）
    for y in [py1 + 10, py2 - 10]:
        draw.line([(px1 - 15, y), (px1, y)], fill=(150, 150, 150), width=1)

    # 底边标注
    draw.text((px1 + rect_w // 2 - 20, py2 + 8), "底边不变", fill=color_label)

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def render_parallelogram_question(doc, q_num: int, q_content: str):
    """
    渲染平行四边形变形题 —— PIL 绘制长方形与平行四边形对比图。

    左：长方形（四角标直角符号），右：拉成的平行四边形（标锐角/钝角弧线），
    中间箭头表示拉伸。直观展示"边长没变，角度变了"。
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    # 嵌入对比图
    img_buf = _draw_parallelogram_comparison()
    if img_buf:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_buf, width=Cm(14))

    # 标注
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_label.paragraph_format.space_after = Pt(2)
    run_label = p_label.add_run("（    ）变了，（    ）没变。")
    set_cjk_font(run_label, FONT, 10)

    doc.add_paragraph()


def render_angle_drawing(doc, q_num: int, q_content: str):
    """渲染画角题 — 留空白让学生画指定的角"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    # 留3行空白用于画角
    for _ in range(3):
        blank_p = doc.add_paragraph()
        blank_p.paragraph_format.space_before = Pt(0)
        blank_p.paragraph_format.space_after = Pt(0)
        run = blank_p.add_run(" " * 60)
        set_cjk_font(run, FONT, BODY_SIZE)


# ============================================================
# 图形符号映射
# ============================================================

# 基础形状
SHAPES = {
    "直角": "∟",
    "锐角": "∠",
    "钝角": "⦥",      # 可用 > 替代
    "三角形": "△",
    "正方形": "□",
    "长方形": "▭",
    "平行四边形": "▱",
    "圆形": "○",
    "五边形": "⬠",
    "六边形": "⬡",
    "交叉": "╳",
    "点": "●",
}

# 角的图标（用于识别题）
ANGLE_ICONS = {
    "直角": "┌",
    "锐角": "∠",
    "钝角": "╲",
}

# 多边形用于数角
POLYGONS_FOR_COUNT = [
    ("△\n三角形", 3),
    ("□\n正方形", 4),
    ("▭\n长方形", 4),
    ("▱\n平行四边形", 4),
    ("⬠\n五边形", 5),
]


# ============================================================
# 钟面图形渲染
# ============================================================

def _draw_clock_image(hour: int, minute: int, second: int = 0,
                       draw_hands: bool = True, size: int = 200) -> BytesIO:
    """
    用 PIL 绘制圆形钟面，返回 PNG 的 BytesIO。

    参数：
        hour: 小时 (0-12)
        minute: 分钟 (0-59)
        second: 秒 (0-59, 默认 0)
        draw_hands: 是否绘制指针（False=空白钟面供学生画）
        size: 图片像素尺寸（默认 200）
    """
    cx = cy = size // 2
    r = int(size * 0.45)

    img = Image.new('RGB', (size, size), 'white')
    draw = ImageDraw.Draw(img)

    # 1. 外圆
    draw.ellipse([cx - r, cy - r, cx + r, cy + r], outline='#333', width=3)

    # 2. 刻度线（60根，正点加粗加长）
    for i in range(60):
        angle_rad = (i * 6 - 90) * math.pi / 180
        is_hour = (i % 5 == 0)
        outer = r - 3
        inner = outer - (12 if is_hour else 6)
        x1 = cx + inner * math.cos(angle_rad)
        y1 = cy + inner * math.sin(angle_rad)
        x2 = cx + outer * math.cos(angle_rad)
        y2 = cy + outer * math.sin(angle_rad)
        draw.line([(x1, y1), (x2, y2)], fill='#333', width=2 if is_hour else 1)

    # 3. 数字 1-12
    for num in range(1, 13):
        angle_rad = (num * 30 - 90) * math.pi / 180
        nr = r - 22
        nx = cx + nr * math.cos(angle_rad)
        ny = cy + nr * math.sin(angle_rad)
        # 用矩形近似文字位置（Pillow 默认字体太小，用较大偏移）
        draw.text((nx - 5, ny - 6), str(num), fill='#333')

    if draw_hands:
        # 4. 时针（短粗，圆头）
        h_angle = ((hour % 12) * 30 + minute * 0.5 - 90) * math.pi / 180
        hx = cx + 38 * math.cos(h_angle)
        hy = cy + 38 * math.sin(h_angle)
        draw.line([(cx, cy), (hx, hy)], fill='#1a1a1a', width=4)

        # 5. 分针（细长，尖头）
        m_angle = (minute * 6 + second * 0.1 - 90) * math.pi / 180
        mx = cx + 58 * math.cos(m_angle)
        my = cy + 58 * math.sin(m_angle)
        draw.line([(cx, cy), (mx, my)], fill='#1a1a1a', width=2)

        # 6. 秒针（最细最长，红色，带尾针）
        s_angle = (second * 6 - 90) * math.pi / 180
        sx = cx + 68 * math.cos(s_angle)
        sy = cy + 68 * math.sin(s_angle)
        # 尾针（反向延伸 15px）
        tx = cx - 15 * math.cos(s_angle)
        ty = cy - 15 * math.sin(s_angle)
        draw.line([(tx, ty), (sx, sy)], fill='#cc0000', width=1)

        # 7. 中心帽（覆盖三针交汇点）
        draw.ellipse([cx - 5, cy - 5, cx + 5, cy + 5], fill='#222')

        # 8. 秒针尾端小红点
        draw.ellipse([cx - 2, cy - 2, cx + 2, cy + 2], fill='#cc0000')

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def render_clock_question(doc, q_num: int, q_content: str, clocks: list):
    """
    渲染钟面读时题。

    clocks: [{"hour": 3, "minute": 0}, ...]
    每个钟面用 PIL 绘制圆形钟面图片，嵌入 Word。
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    cols = min(len(clocks), 4)
    outer = doc.add_table(rows=2, cols=cols)
    outer.alignment = 1
    set_table_borders(outer, val="none")

    for ci, clock in enumerate(clocks):
        h = clock.get("hour", 0) % 12
        if h == 0:
            h = 12
        m = clock.get("minute", 0)
        s = clock.get("second", 0)

        # 图片行
        img_cell = outer.cell(0, ci)
        set_cell_border(img_cell, top="none", bottom="none", left="none", right="none")
        img_cell.width = Cm(3.0)
        img_buf = _draw_clock_image(h, m, s, draw_hands=True)
        p_img = img_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after = Pt(2)
        run_img = p_img.add_run()
        run_img.add_picture(img_buf, width=Cm(2.5))

        # 答案行
        ans_cell = outer.cell(1, ci)
        set_cell_border(ans_cell, top="none", bottom="none", left="none", right="none")
        p_ans = ans_cell.paragraphs[0]
        p_ans.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ans.paragraph_format.space_before = Pt(2)
        p_ans.paragraph_format.space_after = Pt(4)
        run_ans = p_ans.add_run(f"钟面{ci + 1}: __时__分")
        set_cjk_font(run_ans, FONT, 9)

    doc.add_paragraph()


def render_clock_time_question(doc, q_num: int, q_content: str,
                               times: list):
    """
    渲染钟面时间题——给出文字+空白钟面，让学生画时针分针。

    times: [("3:00", "3时"), ("7:30", "7时30分"), ...]
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    cols = min(len(times), 4)
    table = doc.add_table(rows=2, cols=cols)
    table.alignment = 1
    set_table_borders(table, val="single", sz="4")

    for ci, (time_str, label) in enumerate(times):
        # 时间标签行
        cell_label = table.cell(0, ci)
        cell_label.width = Cm(3.0)
        pl = cell_label.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = pl.add_run(f"{label}\n({time_str})")
        set_cjk_font(run_l, FONT, 9)

        # 空白钟面图片
        cell_draw = table.cell(1, ci)
        cell_draw.width = Cm(3.0)
        pd = cell_draw.paragraphs[0]
        pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pd.paragraph_format.space_before = Pt(4)
        pd.paragraph_format.space_after = Pt(4)
        img_buf = _draw_clock_image(0, 0, 0, draw_hands=False)
        run_d = pd.add_run()
        run_d.add_picture(img_buf, width=Cm(2.5))

    doc.add_paragraph()


# ============================================================
# 立方体堆叠渲染
# ============================================================

def _draw_isometric_cubes(grid: list, a: int = 36, b: int = 18,
                          h: int = 32) -> BytesIO:
    """
    用 PIL 绘制等轴测立体立方体堆叠图。

    参数：
        grid: 二维数组，grid[r][c] = 该位置的立方体层数
        a: 顶部菱形半宽（像素）
        b: 顶部菱形半高（像素）
        h: 立方体高度（像素）

    三个可见面配色：
        顶面 — 最亮，左侧面 — 中等，右侧面 — 最暗
    绘制顺序从远到近、从下到上，保证遮挡正确。
    """
    rows = len(grid)
    cols = len(grid[0]) if grid else 0
    if rows == 0 or cols == 0:
        return BytesIO()

    max_layers = max(max(row) for row in grid)

    # 计算图片尺寸
    margin = 40
    img_w = (rows + cols) * a + margin * 2
    img_h = max_layers * h + (rows + cols) * b + h + margin * 2

    origin_x = img_w // 2
    origin_y = margin + max_layers * h + b

    img = Image.new('RGB', (img_w, img_h), 'white')
    draw = ImageDraw.Draw(img)

    # 收集所有需要绘制的立方体，按深度排序
    cubes = []
    for r in range(rows):
        for c in range(cols):
            for l in range(grid[r][c]):
                cubes.append((r, c, l))

    # 排序：低层先画，(row+col) 小的（远的）先画
    cubes.sort(key=lambda x: (x[2], x[0] + x[1]))

    # 三种面的颜色（打印友好的灰度梯度）
    color_top = (230, 238, 250)     # 顶面 — 最亮
    color_left = (170, 195, 225)    # 左侧面 — 中等
    color_right = (125, 155, 195)   # 右侧面 — 最暗
    color_outline = (70, 70, 70)    # 边线

    for r, c, l in cubes:
        # 顶部菱形中心
        cx = origin_x + (c - r) * a
        cy = origin_y + (r + c) * b - l * h

        # 顶部菱形四个顶点
        top_pt = (cx, cy - b)           # 上
        right_pt = (cx + a, cy)         # 右
        bottom_pt = (cx, cy + b)        # 下
        left_pt = (cx - a, cy)          # 左

        # 左侧面：左→下→下+h→左+h
        left_face = [left_pt, bottom_pt,
                     (cx, cy + b + h), (cx - a, cy + h)]
        # 右侧面：右→下→下+h→右+h
        right_face = [right_pt, bottom_pt,
                      (cx, cy + b + h), (cx + a, cy + h)]
        # 顶面菱形
        top_face = [top_pt, right_pt, bottom_pt, left_pt]

        # 绘制面
        draw.polygon(left_face, fill=color_left, outline=color_outline)
        draw.polygon(right_face, fill=color_right, outline=color_outline)
        draw.polygon(top_face, fill=color_top, outline=color_outline)

        # 顶面十字线（增强立体感）
        draw.line([top_pt, bottom_pt], fill=color_outline, width=1)
        draw.line([left_pt, right_pt], fill=color_outline, width=1)

    # ---- 视角文字标注 ----
    dir_font = _get_pil_font(14)
    dir_text = "从右前方看"
    tx = img_w - 110
    ty = img_h - 25
    draw.text((tx, ty), dir_text, fill=(60, 60, 60), font=dir_font)

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def render_cube_stack_question(doc, q_num: int, q_content: str,
                               grid: list, answer: str = ""):
    """
    渲染立方体堆叠计数题 —— 使用等轴测立体图。

    grid: 二维数组，grid[r][c] = 该位置的立方体层数
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    # 生成等轴测立体图并嵌入 Word
    img_buf = _draw_isometric_cubes(grid)
    if img_buf:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(6)
        run_img = p_img.add_run()
        run_img.add_picture(img_buf, width=Cm(7))

    # 问题作答区
    p_ans = doc.add_paragraph()
    p_ans.paragraph_format.space_before = Pt(4)
    run_ans = p_ans.add_run(f"一共有（    ）个小立方体。")
    set_cjk_font(run_ans, FONT, BODY_SIZE)

    doc.add_paragraph()


def _draw_three_view_diagram(front_view: list, side_view: list,
                            top_view: list = None) -> BytesIO:
    """
    用 PIL 绘制三视图立方体示意图。

    每个视图是填充网格：有色格=有立方体，白色格=空。
    三视图并排，清晰标注。
    """
    cell = 40          # 每格像素
    lw = 2             # 线宽
    pad = 50           # 图间距
    margin = 30
    label_h = 30       # 标签高度

    max_h_front = max(front_view) if front_view else 3
    cols_front = len(front_view)
    max_h_side = max(side_view) if side_view else 3
    cols_side = len(side_view)

    # 正面视图尺寸
    fw = cols_front * cell + lw
    fh = max_h_front * cell + lw + label_h

    # 侧面视图尺寸
    sw = cols_side * cell + lw
    sh = max_h_side * cell + lw + label_h

    # 俯视图（如有）
    has_top = top_view and len(top_view) > 0 and isinstance(top_view[0], list)
    if has_top:
        top_rows = len(top_view)
        top_cols = len(top_view[0])
        tw = top_cols * cell + lw
        th = top_rows * cell + lw + label_h
    else:
        tw = th = 0

    label_space = 25  # 底部方向标签空间
    img_w = fw + pad + sw + (pad + tw if has_top else 0) + margin * 2
    img_h = max(fh, sh, th) + margin * 2 + label_space

    img = Image.new('RGB', (img_w, img_h), 'white')
    draw = ImageDraw.Draw(img)

    color_fill = (180, 200, 230)    # 填充色
    color_grid = (50, 50, 50)      # 网格线
    color_label = (60, 60, 60)     # 标签色
    view_font = _get_pil_font(14)  # 方向标注字体

    def draw_grid_view(ox: int, oy: int, cols: int, max_h: int,
                       heights: list, label: str, arrow_from: str = "front"):
        """在 (ox, oy) 处绘制一个柱状视图，下方大号文字标注方向"""
        for ci in range(cols):
            h = heights[ci] if ci < len(heights) else 0
            for ri in range(h):
                rx = ox + ci * cell + lw
                ry = oy + (max_h - 1 - ri) * cell + lw
                draw.rectangle(
                    [rx, ry, rx + cell - lw, ry + cell - lw],
                    fill=color_fill, outline=color_grid, width=1
                )
            for ri in range(h, max_h):
                rx = ox + ci * cell + lw
                ry = oy + (max_h - 1 - ri) * cell + lw
                draw.rectangle(
                    [rx, ry, rx + cell - lw, ry + cell - lw],
                    outline=(200, 200, 200), width=1
                )

        gw = cols * cell
        gh = max_h * cell
        draw.rectangle([ox, oy, ox + gw, oy + gh], outline=color_grid, width=lw)

        # 大号文字标注方向
        lbl_y = oy + gh + 5
        draw.text((ox + gw // 2 - 30, lbl_y), label, fill=color_label, font=view_font)

    def draw_top_grid(ox: int, oy: int, grid_data: list, label: str):
        """绘制俯视图网格，下方大号文字标注方向"""
        rows = len(grid_data)
        cols = len(grid_data[0]) if rows else 0
        for r in range(rows):
            for c in range(cols):
                rx = ox + c * cell + lw
                ry = oy + r * cell + lw
                val = grid_data[r][c] if r < len(grid_data) and c < len(grid_data[r]) else 0
                if val and val > 0:
                    draw.rectangle(
                        [rx, ry, rx + cell - lw, ry + cell - lw],
                        fill=color_fill, outline=color_grid, width=1
                    )
                    draw.text((rx + cell//3, ry + cell//4), str(val),
                              fill=(80, 80, 80))
                else:
                    draw.rectangle(
                        [rx, ry, rx + cell - lw, ry + cell - lw],
                        outline=(200, 200, 200), width=1
                    )

        gw = cols * cell
        gh = rows * cell
        draw.rectangle([ox, oy, ox + gw, oy + gh], outline=color_grid, width=lw)

        # 大号文字标注方向
        lbl_y = oy + gh + 5
        draw.text((ox + gw // 2 - 30, lbl_y), label, fill=color_label, font=view_font)

    # 放置三个视图
    x1 = margin
    y1 = margin
    draw_grid_view(x1, y1, cols_front, max_h_front, front_view, "从正面看", arrow_from="front")

    x2 = x1 + fw + pad
    draw_grid_view(x2, y1, cols_side, max_h_side, side_view, "从侧面看", arrow_from="side")

    if has_top:
        x3 = x2 + sw + pad
        draw_top_grid(x3, y1, top_view, "从上面看")

    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def render_cube_view_question(doc, q_num: int, q_content: str,
                              front_view: list, side_view: list,
                              top_view: list = None):
    """
    渲染三视图立方体计数题 —— PIL 绘制清晰三视图网格。

    front_view: 正面每列层数，如 [3, 2, 1]
    side_view: 侧面每列层数，如 [2, 1, 1]
    top_view: 俯视图二维网格（可选）
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{q_num}. {q_content}")
    set_cjk_font(run, FONT, BODY_SIZE)

    # 生成三视图并嵌入
    img_buf = _draw_three_view_diagram(front_view, side_view, top_view)
    if img_buf:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_buf, width=Cm(12))

    doc.add_paragraph()
    p_ans = doc.add_paragraph()
    run_ans = p_ans.add_run("一共有（    ）个小立方体。")
    set_cjk_font(run_ans, FONT, BODY_SIZE)
    doc.add_paragraph()
