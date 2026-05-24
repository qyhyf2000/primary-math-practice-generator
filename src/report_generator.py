"""学习报告生成器 — 生成周报/月报 docx"""
import sqlite3
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os
sys_path = Path(__file__).parent.parent
import sys
sys.path.insert(0, str(sys_path))


def generate_weekly_report(db_conn: sqlite3.Connection, output_path: str, weeks: int = 1) -> str:
    """生成学习周报 docx"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"小学数学练习 — 学习周报")
    run.font.size = Pt(18)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_date = datetime.now()
    start_date = end_date - timedelta(days=weeks * 7)
    run2 = p2.add_run(f"{start_date.strftime('%Y-%m-%d')} ~ {end_date.strftime('%Y-%m-%d')}")
    run2.font.size = Pt(11)

    doc.add_paragraph()

    # 试卷完成统计
    exam_count = db_conn.execute("""
        SELECT COUNT(DISTINCT exam_title) FROM exam_history
        WHERE used_at >= ?
    """, (start_date.strftime('%Y-%m-%d'),)).fetchone()[0]

    total_qs = db_conn.execute("""
        SELECT COUNT(*) FROM exam_history
        WHERE used_at >= ?
    """, (start_date.strftime('%Y-%m-%d'),)).fetchone()[0]

    h3 = doc.add_paragraph()
    run3 = h3.add_run(f"本周完成: {exam_count} 份试卷，共 {total_qs} 题")
    run3.font.size = Pt(12)
    run3.bold = True

    # 知识点分布
    doc.add_paragraph()
    h4 = doc.add_paragraph()
    run4 = h4.add_run("知识点覆盖:")
    run4.font.size = Pt(12)
    run4.bold = True

    kp_rows = db_conn.execute("""
        SELECT q.knowledge_point, COUNT(*) as cnt
        FROM exam_history eh
        JOIN questions q ON eh.question_id = q.id
        WHERE eh.used_at >= ?
        GROUP BY q.knowledge_point
        ORDER BY cnt DESC
        LIMIT 15
    """, (start_date.strftime('%Y-%m-%d'),)).fetchall()

    if kp_rows:
        table = doc.add_table(rows=len(kp_rows) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.cell(0, 0).text = "知识点"
        table.cell(0, 1).text = "练习次数"
        for i, (kp, cnt) in enumerate(kp_rows):
            table.cell(i + 1, 0).text = kp
            table.cell(i + 1, 1).text = str(cnt)

    # 错题统计
    doc.add_paragraph()
    h5 = doc.add_paragraph()
    run5 = h5.add_run("错题本状态:")
    run5.font.size = Pt(12)
    run5.bold = True

    try:
        wrong_total = db_conn.execute(
            "SELECT COUNT(DISTINCT question_id) FROM wrong_answers"
        ).fetchone()[0]
        wrong_new = db_conn.execute(
            "SELECT COUNT(DISTINCT question_id) FROM wrong_answers WHERE reviewed=0"
        ).fetchone()[0]
        week_wrong = db_conn.execute(
            "SELECT COUNT(DISTINCT question_id) FROM wrong_answers WHERE wrong_at >= ?",
            (start_date.strftime('%Y-%m-%d'),)
        ).fetchone()[0]
        p_w = doc.add_paragraph()
        p_w.add_run(f"错题总数: {wrong_total} | 未复习: {wrong_new} | 本周新增: {week_wrong}").font.size = Pt(11)
    except sqlite3.OperationalError:
        p_w2 = doc.add_paragraph()
        p_w2.add_run("错题本尚未启用").font.size = Pt(11)

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    return output_path
